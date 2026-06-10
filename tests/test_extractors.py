"""Tests for distillcore.extractors."""

from pathlib import Path
from unittest.mock import MagicMock, patch

import pytest

from distillcore.extractors import (
    _detect_format,
    _validate_path,
    extract,
    get_registered_formats,
    register_extractor,
)
from distillcore.extractors.csv import CsvExtractor
from distillcore.extractors.text import TextExtractor
from distillcore.models import ExtractionResult, PageText


class TestTextExtractor:
    def test_extract_txt(self, tmp_path: Path) -> None:
        f = tmp_path / "test.txt"
        f.write_text("Hello world\n\nSecond paragraph")
        extractor = TextExtractor()
        result = extractor.extract(f)
        assert result.format == "txt"
        assert result.page_count == 1
        assert "Hello world" in result.full_text
        assert result.pages[0].page_number == 1

    def test_extract_markdown(self, tmp_path: Path) -> None:
        f = tmp_path / "doc.md"
        f.write_text("# Title\n\nContent here")
        extractor = TextExtractor()
        result = extractor.extract(f)
        assert "# Title" in result.full_text

    def test_formats(self) -> None:
        assert "txt" in TextExtractor.formats
        assert "md" in TextExtractor.formats


class TestDetectFormat:
    def test_pdf(self) -> None:
        assert _detect_format(Path("file.pdf")) == "pdf"

    def test_txt(self) -> None:
        assert _detect_format(Path("file.txt")) == "txt"

    def test_no_extension(self) -> None:
        assert _detect_format(Path("README")) == "txt"

    def test_uppercase(self) -> None:
        assert _detect_format(Path("FILE.PDF")) == "pdf"


class TestPathValidation:
    def test_unrestricted(self, tmp_path: Path) -> None:
        f = tmp_path / "test.txt"
        f.write_text("data")
        result = _validate_path(f, allowed_dirs=None)
        assert result == f.resolve()

    def test_within_allowed_dir(self, tmp_path: Path) -> None:
        f = tmp_path / "test.txt"
        f.write_text("data")
        result = _validate_path(f, allowed_dirs=[str(tmp_path)])
        assert result == f.resolve()

    def test_outside_allowed_dir(self, tmp_path: Path) -> None:
        with pytest.raises(PermissionError, match="Access denied"):
            _validate_path(Path("/etc/passwd"), allowed_dirs=[str(tmp_path)])

    def test_traversal_rejected(self, tmp_path: Path) -> None:
        evil = tmp_path / ".." / ".." / "etc" / "passwd"
        with pytest.raises(PermissionError, match="Access denied"):
            _validate_path(evil, allowed_dirs=[str(tmp_path)])

    def test_multiple_allowed_dirs(self, tmp_path: Path) -> None:
        subdir = tmp_path / "docs"
        subdir.mkdir()
        f = subdir / "test.txt"
        f.write_text("data")
        result = _validate_path(f, allowed_dirs=["/nonexistent", str(tmp_path)])
        assert result == f.resolve()


class TestExtractRegistry:
    def test_text_registered(self) -> None:
        assert "txt" in get_registered_formats()
        assert "md" in get_registered_formats()

    def test_extract_text_file(self, tmp_path: Path) -> None:
        f = tmp_path / "test.txt"
        f.write_text("content")
        result = extract(f)
        assert result.full_text == "content"

    def test_extract_with_format_override(self, tmp_path: Path) -> None:
        f = tmp_path / "data.dat"
        f.write_text("plain text data")
        result = extract(f, format="txt")
        assert result.full_text == "plain text data"

    def test_unknown_format_raises(self, tmp_path: Path) -> None:
        f = tmp_path / "file.xyz"
        f.write_text("data")
        with pytest.raises(ValueError, match="No extractor registered"):
            extract(f)

    def test_custom_extractor(self, tmp_path: Path) -> None:
        # Use a clearly-fake extension so this test doesn't override the
        # built-in CsvExtractor registered for "csv".
        class FakeExtractor:
            formats = ["fakefmt"]

            def extract(self, source, config=None):
                return ExtractionResult(
                    pages=[PageText(page_number=1, text="fake data")],
                    full_text="fake data",
                    page_count=1,
                    format="fakefmt",
                )

        register_extractor(FakeExtractor())
        f = tmp_path / "data.fakefmt"
        f.write_text("a,b,c")
        result = extract(f)
        assert result.format == "fakefmt"

    def test_text_extractor_ignores_config(self, tmp_path: Path) -> None:
        f = tmp_path / "test.txt"
        f.write_text("content")
        extractor = TextExtractor()
        result = extractor.extract(f, config={"arbitrary": "config"})
        assert result.full_text == "content"


class TestDocxExtractor:
    def test_extract_paragraphs(self, tmp_path: Path) -> None:
        """Create a real .docx and extract it."""
        import docx

        doc = docx.Document()
        doc.core_properties.title = "Test Report"
        doc.core_properties.author = "Jane Doe"
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        f = tmp_path / "test.docx"
        doc.save(str(f))

        from distillcore.extractors.docx import DocxExtractor

        extractor = DocxExtractor()
        result = extractor.extract(f)

        assert result.format == "docx"
        assert result.page_count == 1
        assert "First paragraph." in result.full_text
        assert "Second paragraph." in result.full_text
        assert "\n\n" in result.full_text
        assert result.metadata["title"] == "Test Report"
        assert result.metadata["author"] == "Jane Doe"

    def test_extract_table(self, tmp_path: Path) -> None:
        """Tables should be extracted as tab-separated rows."""
        import docx

        doc = docx.Document()
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "A1"
        table.cell(0, 1).text = "B1"
        table.cell(1, 0).text = "A2"
        table.cell(1, 1).text = "B2"
        f = tmp_path / "table.docx"
        doc.save(str(f))

        from distillcore.extractors.docx import DocxExtractor

        result = DocxExtractor().extract(f)
        assert "A1\tB1" in result.full_text
        assert "A2\tB2" in result.full_text

    def test_registry_detection(self, tmp_path: Path) -> None:
        assert "docx" in get_registered_formats()

    def test_empty_doc(self, tmp_path: Path) -> None:
        import docx

        doc = docx.Document()
        f = tmp_path / "empty.docx"
        doc.save(str(f))

        from distillcore.extractors.docx import DocxExtractor

        result = DocxExtractor().extract(f)
        assert result.full_text == ""
        assert result.page_count == 1


class TestHtmlExtractor:
    def test_basic_extraction(self, tmp_path: Path) -> None:
        f = tmp_path / "page.html"
        f.write_text(
            "<html><head><title>My Page</title></head>"
            "<body><p>First paragraph.</p><p>Second paragraph.</p></body></html>"
        )
        from distillcore.extractors.html import HtmlExtractor

        result = HtmlExtractor().extract(f)
        assert result.format == "html"
        assert "First paragraph." in result.full_text
        assert "Second paragraph." in result.full_text
        assert result.metadata["title"] == "My Page"

    def test_strips_script_and_style(self, tmp_path: Path) -> None:
        f = tmp_path / "noisy.html"
        f.write_text(
            "<html><body>"
            "<script>var x = 1;</script>"
            "<style>body { color: red; }</style>"
            "<p>Real content.</p>"
            "</body></html>"
        )
        from distillcore.extractors.html import HtmlExtractor

        result = HtmlExtractor().extract(f)
        assert "var x" not in result.full_text
        assert "color: red" not in result.full_text
        assert "Real content." in result.full_text

    def test_strips_nav_footer(self, tmp_path: Path) -> None:
        f = tmp_path / "layout.html"
        f.write_text(
            "<html><body>"
            "<nav>Menu items</nav>"
            "<main><p>Main content.</p></main>"
            "<footer>Copyright 2026</footer>"
            "</body></html>"
        )
        from distillcore.extractors.html import HtmlExtractor

        result = HtmlExtractor().extract(f)
        assert "Menu items" not in result.full_text
        assert "Copyright" not in result.full_text
        assert "Main content." in result.full_text

    def test_extracts_author_meta(self, tmp_path: Path) -> None:
        f = tmp_path / "meta.html"
        f.write_text(
            '<html><head><meta name="author" content="Jane Doe"></head>'
            "<body><p>Content.</p></body></html>"
        )
        from distillcore.extractors.html import HtmlExtractor

        result = HtmlExtractor().extract(f)
        assert result.metadata["author"] == "Jane Doe"

    def test_empty_html(self, tmp_path: Path) -> None:
        f = tmp_path / "empty.html"
        f.write_text("<html><body></body></html>")
        from distillcore.extractors.html import HtmlExtractor

        result = HtmlExtractor().extract(f)
        assert result.page_count == 1

    def test_registry_detection(self) -> None:
        formats = get_registered_formats()
        assert "html" in formats
        assert "htm" in formats


class TestPdfExtractor:
    def test_extract_pdf(self, tmp_path: Path) -> None:
        """Test PDF extraction with mocked pdfplumber."""
        mock_page = MagicMock()
        mock_page.extract_text.return_value = "Page one text"

        mock_pdf = MagicMock()
        mock_pdf.__enter__ = MagicMock(return_value=mock_pdf)
        mock_pdf.__exit__ = MagicMock(return_value=False)
        mock_pdf.pages = [mock_page]

        with patch("distillcore.extractors.pdf.pdfplumber") as mock_plumber:
            mock_plumber.open.return_value = mock_pdf
            from distillcore.extractors.pdf import PdfExtractor

            extractor = PdfExtractor()
            result = extractor.extract(tmp_path / "test.pdf")

        assert result.format == "pdf"
        assert result.page_count == 1
        assert "Page one text" in result.full_text


class TestCsvExtractor:
    def test_basic_csv(self, tmp_path: Path) -> None:
        f = tmp_path / "data.csv"
        f.write_text("name,age,city\nAlice,30,NYC\nBob,25,LA")
        result = CsvExtractor().extract(f)
        assert result.format == "csv"
        assert result.page_count == 1
        # Tab-normalized in the output regardless of source delimiter
        assert "name\tage\tcity" in result.full_text
        assert "Alice\t30\tNYC" in result.full_text
        assert result.metadata["columns"] == ["name", "age", "city"]
        assert result.metadata["row_count"] == 2
        assert result.metadata["delimiter"] == ","

    def test_tsv(self, tmp_path: Path) -> None:
        f = tmp_path / "data.tsv"
        f.write_text("col1\tcol2\nv1\tv2\nv3\tv4")
        result = CsvExtractor().extract(f)
        assert result.format == "tsv"
        assert result.metadata["delimiter"] == "\t"
        assert result.metadata["columns"] == ["col1", "col2"]
        assert result.metadata["row_count"] == 2

    def test_pipe_separated(self, tmp_path: Path) -> None:
        f = tmp_path / "data.csv"
        f.write_text("a|b|c\n1|2|3\n4|5|6")
        result = CsvExtractor().extract(f)
        assert result.metadata["delimiter"] == "|"
        # Output is tab-normalized
        assert "a\tb\tc" in result.full_text
        assert "1\t2\t3" in result.full_text

    def test_quoted_cells_with_commas(self, tmp_path: Path) -> None:
        f = tmp_path / "data.csv"
        f.write_text('name,note\nAlice,"hello, world"\nBob,"a,b,c"')
        result = CsvExtractor().extract(f)
        assert result.metadata["row_count"] == 2
        # Embedded commas preserved in cells, not split
        assert "hello, world" in result.full_text
        assert "a,b,c" in result.full_text

    def test_quoted_cells_with_newlines(self, tmp_path: Path) -> None:
        # csv.reader handles quoted newlines when given the file directly; via
        # splitlines() they become two rows. This documents that limitation.
        f = tmp_path / "data.csv"
        f.write_text('a,b\n"line1\nline2",x')
        result = CsvExtractor().extract(f)
        # Just verify we got *something* without crashing
        assert result.format == "csv"
        assert result.metadata["row_count"] >= 1

    def test_empty_file(self, tmp_path: Path) -> None:
        f = tmp_path / "empty.csv"
        f.write_text("")
        result = CsvExtractor().extract(f)
        assert result.page_count == 0
        assert result.pages == []
        assert result.metadata["columns"] == []
        assert result.metadata["row_count"] == 0

    def test_whitespace_only(self, tmp_path: Path) -> None:
        f = tmp_path / "blank.csv"
        f.write_text("   \n  \n")
        result = CsvExtractor().extract(f)
        assert result.page_count == 0

    def test_header_only(self, tmp_path: Path) -> None:
        f = tmp_path / "headeronly.csv"
        f.write_text("a,b,c")
        result = CsvExtractor().extract(f)
        assert result.metadata["columns"] == ["a", "b", "c"]
        assert result.metadata["row_count"] == 0
        # Header still appears in the output
        assert "a\tb\tc" in result.full_text

    def test_single_column(self, tmp_path: Path) -> None:
        # Sniffer can't detect delimiter on single-column files; falls back to ext.
        f = tmp_path / "single.csv"
        f.write_text("name\nAlice\nBob")
        result = CsvExtractor().extract(f)
        assert result.metadata["row_count"] == 2

    def test_via_registry(self, tmp_path: Path) -> None:
        """End-to-end via the top-level extract() entry point."""
        f = tmp_path / "data.csv"
        f.write_text("a,b\n1,2")
        result = extract(f)
        assert result.format == "csv"
        assert result.metadata["columns"] == ["a", "b"]

    def test_formats(self) -> None:
        assert "csv" in CsvExtractor.formats
        assert "tsv" in CsvExtractor.formats


class TestExcelExtractor:
    """Build real .xlsx files and round-trip them through the extractor."""

    @staticmethod
    def _make_xlsx(path, sheets):
        """sheets: dict of {sheet_name: [[row1_cells], [row2_cells], ...]}."""
        import openpyxl

        wb = openpyxl.Workbook()
        # Remove default sheet; we'll add ours by name
        wb.remove(wb.active)
        for name, rows in sheets.items():
            ws = wb.create_sheet(title=name)
            for row in rows:
                ws.append(row)
        wb.save(str(path))

    def test_single_sheet(self, tmp_path: Path) -> None:
        f = tmp_path / "single.xlsx"
        self._make_xlsx(f, {"Sheet1": [["name", "age"], ["Alice", 30], ["Bob", 25]]})
        from distillcore.extractors.excel import ExcelExtractor

        result = ExcelExtractor().extract(f)
        assert result.format == "xlsx"
        assert result.page_count == 1
        assert result.pages[0].page_number == 1
        assert "name\tage" in result.full_text
        assert "Alice\t30" in result.full_text
        assert "Bob\t25" in result.full_text
        assert result.metadata["sheet_names"] == ["Sheet1"]
        assert result.metadata["row_counts"]["Sheet1"] == 3

    def test_multi_sheet_page_numbers(self, tmp_path: Path) -> None:
        f = tmp_path / "multi.xlsx"
        self._make_xlsx(
            f,
            {
                "First": [["a", "b"], [1, 2]],
                "Second": [["c", "d"], [3, 4]],
                "Third": [["e", "f"], [5, 6]],
            },
        )
        from distillcore.extractors.excel import ExcelExtractor

        result = ExcelExtractor().extract(f)
        assert result.page_count == 3
        assert [p.page_number for p in result.pages] == [1, 2, 3]
        # Sheet order preserved
        assert result.metadata["sheet_names"] == ["First", "Second", "Third"]
        assert "a\tb" in result.pages[0].text
        assert "c\td" in result.pages[1].text
        assert "e\tf" in result.pages[2].text

    def test_mixed_cell_types(self, tmp_path: Path) -> None:
        from datetime import date, datetime

        f = tmp_path / "mixed.xlsx"
        self._make_xlsx(
            f,
            {
                "Data": [
                    ["str", "int", "float", "bool", "date", "datetime"],
                    ["hello", 42, 3.14, True, date(2026, 1, 15), datetime(2026, 1, 15, 12, 30)],
                ]
            },
        )
        from distillcore.extractors.excel import ExcelExtractor

        result = ExcelExtractor().extract(f)
        text = result.full_text
        assert "hello" in text
        assert "42" in text
        assert "3.14" in text
        assert "TRUE" in text
        assert "2026-01-15" in text
        # datetime renders as full isoformat
        assert "2026-01-15T12:30:00" in text

    def test_empty_rows_skipped(self, tmp_path: Path) -> None:
        f = tmp_path / "gaps.xlsx"
        self._make_xlsx(
            f,
            {
                "Sheet1": [
                    ["a", "b"],
                    [None, None],  # all-empty row
                    ["1", "2"],
                    [None, None],
                    ["3", "4"],
                ]
            },
        )
        from distillcore.extractors.excel import ExcelExtractor

        result = ExcelExtractor().extract(f)
        # 3 non-empty rows: header + 2 data
        assert result.metadata["row_counts"]["Sheet1"] == 3
        text = result.full_text
        assert "a\tb" in text
        assert "1\t2" in text
        assert "3\t4" in text

    def test_empty_sheet_dropped(self, tmp_path: Path) -> None:
        f = tmp_path / "with_empty.xlsx"
        self._make_xlsx(
            f,
            {
                "WithData": [["a", "b"], [1, 2]],
                "Empty": [],
                "MoreData": [["x", "y"], [3, 4]],
            },
        )
        from distillcore.extractors.excel import ExcelExtractor

        result = ExcelExtractor().extract(f)
        # Empty sheet is in sheet_names + row_counts but not in pages.
        assert result.metadata["sheet_names"] == ["WithData", "Empty", "MoreData"]
        assert result.metadata["row_counts"]["Empty"] == 0
        assert result.page_count == 2
        # Page numbers compact across the skipped sheet
        assert [p.page_number for p in result.pages] == [1, 2]

    def test_via_registry(self, tmp_path: Path) -> None:
        f = tmp_path / "data.xlsx"
        self._make_xlsx(f, {"Sheet1": [["a", "b"], [1, 2]]})
        result = extract(f)
        assert result.format == "xlsx"
        assert result.metadata["sheet_names"] == ["Sheet1"]

    def test_registry_detection(self) -> None:
        assert "xlsx" in get_registered_formats()

    def test_formats(self) -> None:
        from distillcore.extractors.excel import ExcelExtractor

        assert "xlsx" in ExcelExtractor.formats
